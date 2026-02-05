import random
from docx import Document
from rlm.rlm_repl import RLMRepl


def generate_massive_context(num_lines: int = 1_000_000, answer: str = "1298418") -> str:
    print("Generating massive context with 1M lines...")

    # Set of random words to use
    random_words = ["blah", "random", "text", "data", "content", "information", "sample"]

    lines = []
    for _ in range(num_lines):
        num_words = random.randint(3, 8)
        line_words = [random.choice(random_words) for _ in range(num_words)]
        lines.append(" ".join(line_words))

    # Insert the magic number at a random position (somewhere in the middle)
    magic_position = random.randint(400000, 600000)
    lines[magic_position] = f"The magic number is {answer}"

    print(f"Magic number inserted at position {magic_position}")

    return "\n".join(lines)


def read_word_document(file_path: str) -> str:
    """
    读取Word文档内容（包括段落和表格）
    """
    doc = Document(file_path)
    content_parts = []

    # 读取段落内容
    for paragraph in doc.paragraphs:
        content_parts.append(paragraph.text)

    # 读取表格内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                content_parts.append(cell.text)
    full_text = '\n'.join(content_parts)
    # 将不间断空格(\xa0)替换为普通空格
    full_text = full_text.replace('\xa0', ' ')

    return full_text


def main():
    # answer = str(random.randint(1000000, 9999999))
    # context = generate_massive_context(num_lines=1_000_000, answer=answer)
    # 读取word文档
    context = read_word_document(r"C:\Users\admin\Desktop\安渡(UniNXG)REST服务接口规范文档1.4.1\安渡(UniNXG)REST服务接口规范文档1.4.1.docx")
    # 这里读取.env-example的数据
    rlm = RLMRepl(
        enable_logging=True,
        max_iterations=10
    )
    query = "帮我找到这个文档里面关于上传文件的API信息"
    result = rlm.completion(context=context, query=query)
    # print(f"Result: {result}. Expected: {answer}")
    print(f"结果为以下：{result}")


if __name__ == '__main__':
    main()
